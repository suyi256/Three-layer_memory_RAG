"""
Word（.docx）解析：将段落与表格转为「带标题路径」的文本片段，供后续分块。

关键点：
- 通过 Word 样式识别标题（中英样式名），维护 `heading_stack` 表示当前章节路径；
- 表格按行拼接为制表符分隔文本，并继承当前标题路径。
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document


def _heading_level_from_digit_char(ch: str) -> int | None:
    """将样式名中的单个数字字符转为 1..9（支持半角与全角数字）。"""
    if not ch:
        return None
    if ch.isdigit():
        return int(ch)
    o = ord(ch)
    # 全角 ０-９ U+FF10..FF19
    if 0xFF10 <= o <= 0xFF19:
        return o - 0xFF10
    return None


def _extract_heading_level(style_name: str | None) -> int:
    """
    从 Word 段落样式名解析标题层级 1..6。

    支持：
    - 英文：`Heading 1` .. `Heading 6`（样式名中第二段为数字）；
    - 中文：`标题1`、`标题 2`、`标题3` 等（「标题」后可选空白再接数字）。
    无法识别时回退为 1（与 Word 默认「标题」无级别数字时一致）。
    """
    if not style_name:
        return 1
    name = style_name.strip()
    lower = name.lower()

    if lower.startswith("heading"):
        parts = name.split()
        if len(parts) > 1 and parts[1].isdigit():
            return max(1, min(int(parts[1]), 6))

    if "标题" in name:
        # 优先匹配「标题」紧邻的半角/全角数字（常见内置样式名）
        m = re.search(r"标题\s*(\d)", name)
        if m:
            lv = _heading_level_from_digit_char(m.group(1))
            if lv is not None:
                return max(1, min(lv, 6))
        # 兜底：样式名中在「标题」之后出现的第一个数字
        m = re.search(r"标题.*?(\d)", name)
        if m:
            lv = _heading_level_from_digit_char(m.group(1))
            if lv is not None:
                return max(1, min(lv, 6))

    return 1


@dataclass(frozen=True)
class ParsedSegment:
    """来自 Word 的一段连续文本及其标题路径（用于检索展示与分块上下文）。"""

    text: str
    heading_path: str


def _is_heading(style_name: str | None) -> bool:
    """判断段落是否为标题样式（兼容中英文模板）。"""
    if not style_name:
        return False
    s = style_name.lower()
    return s.startswith("heading") or "标题" in style_name


def parse_docx_bytes(data: bytes) -> list[ParsedSegment]:
    """解析 .docx：段落走标题栈；表格单独成段。"""
    doc = Document(io.BytesIO(data))
    heading_stack: list[str] = []
    segments: list[ParsedSegment] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        style_name = para.style.name if para.style is not None else None
        if _is_heading(style_name) and text:
            # 英文 Heading N 与中文「标题N」均需解析层级，否则中文标题会全部视为 1 级，大纲错乱
            level = _extract_heading_level(style_name)
            while len(heading_stack) >= level:
                heading_stack.pop()
            heading_stack.append(text)
            continue
        if not text:
            continue
        path = " / ".join(heading_stack) if heading_stack else ""
        segments.append(ParsedSegment(text=text, heading_path=path))

    for table in doc.tables:
        rows_out: list[str] = []
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            if any(cells):
                rows_out.append("\t".join(cells))
        if rows_out:
            block = "\n".join(rows_out)
            path = " / ".join(heading_stack) if heading_stack else ""
            segments.append(ParsedSegment(text=block, heading_path=path))

    return segments
